import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import pandas as pd

st.set_page_config(page_title="Генератор спецификации", layout="wide")
st.title("📄 Генератор спецификации по программам")

PROGRAM_OPTIONS = [
    "Case.one",
    "Case.one тариф Управляй делами",
    "Doc.one",
    "Bot.one",
    "Casebook тариф Standard",
    "Casebook тариф PRO",
    "Caselook",
    "Casebook API"
]

if "rows" not in st.session_state or len(st.session_state.rows) == 0:
    st.session_state.rows = [{
        "name": PROGRAM_OPTIONS[0],
        "start_date": datetime.today().date(),
        "end_date": datetime.today().date(),
        "count": 1,
        "price_annual": 0.0
    }]

# Заголовки один раз
st.markdown("""
<style>
.field-labels {
    display: flex;
    font-weight: bold;
    padding-left: 10px;
    margin-bottom: 4px;
}
.field-labels > div {
    flex: 1;
}
</style>
<div class="field-labels">
    <div>Наименование программы</div>
    <div>Начало предоставления доступа</div>
    <div>Окончание предоставления доступа</div>
    <div>Кол-во</div>
    <div>₽ за 12 мес</div>
</div>
""", unsafe_allow_html=True)

valid_rows = []
for i, row in enumerate(st.session_state.rows):
    cols = st.columns([1.5, 1, 1, 0.7, 0.8, 0.25, 0.25])
    with cols[0]:
        row["name"] = st.selectbox("", PROGRAM_OPTIONS, key=f"name_{i}")
    with cols[1]:
        row["start_date"] = st.date_input("", value=row["start_date"], format="DD.MM.YYYY", key=f"start_{i}")
    with cols[2]:
        row["end_date"] = st.date_input("", value=row["end_date"], format="DD.MM.YYYY", key=f"end_{i}")
    with cols[3]:
        row["count"] = st.number_input("", min_value=1, step=1, value=row["count"], key=f"count_{i}", label_visibility="collapsed")
    with cols[4]:
        row["price_annual"] = st.number_input("", min_value=0.0, step=100.0, value=row["price_annual"], key=f"price_{i}", format="%.2f", label_visibility="collapsed")

    # Удалить
    with cols[5]:
        if st.button("🗑️", key=f"del_{i}"):
            st.session_state.rows.pop(i)
            st.experimental_rerun()
    # Добавить (только последняя строка)
    with cols[6]:
        if i == len(st.session_state.rows) - 1:
            if st.button("➕", key=f"add_{i}"):
                st.session_state.rows.append({
                    "name": PROGRAM_OPTIONS[0],
                    "start_date": datetime.today().date(),
                    "end_date": datetime.today().date(),
                    "count": 1,
                    "price_annual": 0.0
                })
                st.experimental_rerun()

    # Проверка на корректность дат
    if row["start_date"] > row["end_date"]:
        st.warning(f"⚠️ В строке {i+1} начальная дата позже конечной.", icon="⚠️")
    else:
        valid_rows.append(row)

# Расчёт стоимости
def calculate_price(start_date, end_date, annual_price):
    days = (end_date - start_date).days + 1
    price_per_day = annual_price / 365
    return round(price_per_day * days, 2)

# Генерация DOCX
def generate_specification_docx(data_rows):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    doc.add_paragraph("Спецификация", style='Normal').runs[0].bold = True

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'

    headers = [
        "№",
        "Правообладатель",
        "Наименование программы для ЭВМ, право использования которой предоставляется Лицензиату",
        "Кол-во Лицензий*",
        "Срок, на который предоставляется право",
        "Цена, руб. РФ",
        "Сумма, руб. РФ"
    ]

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), "D9D9D9")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)

    total_sum = 0
    for idx, row in enumerate(data_rows, 1):
        cells = table.add_row().cells
        name = f"Программа для ЭВМ {row['name']}"
        count = row["count"]
        period = f"с {row['start_date'].strftime('%d.%m.%Y')} по {row['end_date'].strftime('%d.%m.%Y')}"
        per_license = row["per_license"]
        total = row["total"]

        values = [
            str(idx),
            'АО "Право.ру"',
            name,
            str(count),
            period,
            f"{per_license:,.2f}".replace(",", " ").replace(".", ","),
            f"{total:,.2f}".replace(",", " ").replace(".", ",")
        ]

        for i, val in enumerate(values):
            cells[i].text = val
            run = cells[i].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)

        total_sum += total

    total_row = table.add_row().cells
    total_row[0].merge(total_row[5])
    total_row[0].text = "Итого общий размер лицензионного вознаграждения:"
    run = total_row[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)

    total_row[6].text = f"{total_sum:,.2f}".replace(",", " ").replace(".", ",")
    run2 = total_row[6].paragraphs[0].runs[0]
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(9)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Таблица и экспорт
if valid_rows:
    data_rows = []
    for row in valid_rows:
        start_dt = datetime.combine(row["start_date"], datetime.min.time())
        end_dt = datetime.combine(row["end_date"], datetime.min.time())
        per_license = calculate_price(start_dt, end_dt, row["price_annual"])
        total = round(per_license * row["count"], 2)
        data_rows.append({
            "name": row["name"],
            "count": row["count"],
            "start_date": row["start_date"],
            "end_date": row["end_date"],
            "per_license": per_license,
            "total": total
        })

    df = pd.DataFrame([{
        "№": idx + 1,
        "Правообладатель": 'АО "Право.ру"',
        "Наименование программы для ЭВМ, право использования которой предоставляется Лицензиату": f"Программа для ЭВМ {r['name']}",
        "Кол-во лицензий": r["count"],
        "Срок": f"от {r['start_date'].strftime('%d.%m.%Y')} до {r['end_date'].strftime('%d.%m.%Y')} гг.",
        "Стоимость лицензии, руб. РФ": f"{r['per_license']:,.2f}".replace(",", " ").replace(".", ","),
        "Сумма, руб. РФ": f"{r['total']:,.2f}".replace(",", " ").replace(".", ",")
    } for idx, r in enumerate(data_rows)])
    
    st.markdown("### 🧾 Расчёт по позициям:")
    st.table(df)

    docx_buffer = generate_specification_docx(data_rows)
    st.download_button(
        label="📥 Скачать спецификацию (.docx)",
        data=docx_buffer,
        file_name="спецификация.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
